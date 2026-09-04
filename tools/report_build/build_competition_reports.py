from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "competition_submission"
FIG_DIR = ROOT / "docs" / "figures"
SCREEN_ROOT = Path(r"D:\矿业\重排截图")


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_center(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fill: str, fnt: ImageFont.ImageFont) -> None:
    lines = text.split("\n")
    line_heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        widths.append(bbox[2] - bbox[0])
        line_heights.append(bbox[3] - bbox[1])
    total_h = sum(line_heights) + (len(lines) - 1) * 8
    y = box[1] + ((box[3] - box[1]) - total_h) / 2
    for line, w, h in zip(lines, widths, line_heights):
        x = box[0] + ((box[2] - box[0]) - w) / 2
        draw.text((x, y), line, fill=fill, font=fnt)
        y += h + 8


def rounded_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fill: str, outline: str = "#5b6b7c") -> None:
    draw.rounded_rectangle(box, radius=22, fill=fill, outline=outline, width=2)
    draw_center(draw, box, text, "#102033", font(26, bold=True))


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#334155") -> None:
    draw.line([start, end], fill=color, width=4)
    x1, y1 = end
    if end[0] >= start[0]:
        pts = [(x1, y1), (x1 - 16, y1 - 9), (x1 - 16, y1 + 9)]
    else:
        pts = [(x1, y1), (x1 + 16, y1 - 9), (x1 + 16, y1 + 9)]
    draw.polygon(pts, fill=color)


def label_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str) -> None:
    x, y = xy
    fnt = font(24, True)
    bbox = draw.textbbox((x, y), text, font=fnt)
    pad_x, pad_y = 10, 6
    rect = (bbox[0] - pad_x, bbox[1] - pad_y, bbox[2] + pad_x, bbox[3] + pad_y)
    draw.rounded_rectangle(rect, radius=10, fill="#f8fafc", outline="#cbd5e1", width=1)
    draw.text((x, y), text, fill="#0f172a", font=fnt)


def create_system_architecture() -> Path:
    path = FIG_DIR / "系统总体架构图.png"
    img = Image.new("RGB", (1800, 1120), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "生产架构：真实数据 - 模型 - 服务 - 多角色协同", fill="#0f172a", font=font(42, True))
    layers = [
        ("感知设备层\n应力/离层/下沉/支架阻力/锚索/微震", "#dbeafe"),
        ("真实数据与治理层\n20,000 行 CSV/校验/卡尔曼平滑/标准化", "#dcfce7"),
        ("XGBoost 模型层\n四级概率/风险分值/特征证据/确定性 JSON", "#fef3c7"),
        ("Render 服务层\nNode API/会话/权限校验/RoofRisk API v1", "#fee2e2"),
        ("Vercel 可视化层\nThree.js/ECharts/三维风险场/静态资源", "#ede9fe"),
        ("Supabase 与多角色应用\nAuth/RBAC/审计/企业/监管/智库/只读/管理员", "#cffafe"),
    ]
    x = 120
    y = 145
    w = 1560
    h = 115
    for i, (label, color) in enumerate(layers):
        box = (x, y + i * 150, x + w, y + i * 150 + h)
        rounded_box(draw, box, label, color)
        if i < len(layers) - 1:
            arrow(draw, (900, box[3] + 8), (900, box[3] + 34))
    img.save(path)
    return path


def create_database_er() -> Path:
    path = FIG_DIR / "数据库ER关系图.png"
    img = Image.new("RGB", (1800, 1050), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "Supabase RBAC 与审计数据关系", fill="#0f172a", font=font(42, True))
    boxes = {
        "auth.users\n认证用户": (90, 170, 430, 310),
        "profiles\n用户资料/状态": (590, 170, 930, 310),
        "user_roles\n用户角色关联": (1090, 170, 1430, 310),
        "roles\n角色": (1090, 470, 1430, 610),
        "role_permissions\n角色权限关联": (590, 470, 930, 610),
        "permissions\n权限": (90, 470, 430, 610),
        "audit_logs\n审计日志": (590, 760, 930, 900),
        "Render API\n服务端校验": (90, 760, 430, 900),
    }
    colors = ["#dbeafe", "#dcfce7", "#fef3c7", "#fee2e2", "#ede9fe", "#cffafe", "#fae8ff", "#e2e8f0"]
    def center(label: str) -> tuple[int, int]:
        b = boxes[label]
        return ((b[0] + b[2]) // 2, (b[1] + b[3]) // 2)
    def edge_points(a: str, b: str) -> tuple[tuple[int, int], tuple[int, int]]:
        ax1, ay1, ax2, ay2 = boxes[a]
        bx1, by1, bx2, by2 = boxes[b]
        ac, bc = center(a), center(b)
        if abs(ac[0] - bc[0]) >= abs(ac[1] - bc[1]):
            start = (ax2 if bc[0] > ac[0] else ax1, ac[1])
            end = (bx1 if bc[0] > ac[0] else bx2, bc[1])
        else:
            start = (ac[0], ay2 if bc[1] > ac[1] else ay1)
            end = (bc[0], by1 if bc[1] > ac[1] else by2)
        return start, end
    relations = [
        ("auth.users\n认证用户", "profiles\n用户资料/状态"),
        ("profiles\n用户资料/状态", "user_roles\n用户角色关联"),
        ("user_roles\n用户角色关联", "roles\n角色"),
        ("roles\n角色", "role_permissions\n角色权限关联"),
        ("role_permissions\n角色权限关联", "permissions\n权限"),
        ("profiles\n用户资料/状态", "audit_logs\n审计日志"),
        ("Render API\n服务端校验", "permissions\n权限"),
    ]
    for a, b in relations:
        start, end = edge_points(a, b)
        arrow(draw, start, end, "#475569")
    for (label, box), color in zip(boxes.items(), colors):
        rounded_box(draw, box, label, color)
    img.save(path)
    return path


def create_monitor_network() -> Path:
    path = FIG_DIR / "监测网络布设示意图.png"
    img = Image.new("RGB", (1800, 1050), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "采掘工作面顶板监测网络布设示意", fill="#0f172a", font=font(42, True))
    draw.rounded_rectangle((170, 210, 1630, 780), radius=26, fill="#e5e7eb", outline="#64748b", width=4)
    draw.rectangle((260, 310, 1540, 680), fill="#334155")
    draw.rectangle((560, 350, 1250, 640), fill="#111827")
    draw.text((615, 460), "1206 工作面 / 采煤机 / 液压支架组", fill="#f8fafc", font=font(34, True))
    draw.text((290, 245), "运输顺槽", fill="#0f172a", font=font(30, True))
    draw.text((1320, 245), "回风顺槽", fill="#0f172a", font=font(30, True))
    sensors = [
        ("应力 S1", (390, 360), "#ef4444"),
        ("应力 S2", (690, 300), "#ef4444"),
        ("离层 D1", (820, 275), "#f59e0b"),
        ("下沉 U1", (980, 295), "#eab308"),
        ("支架 P1", (710, 620), "#3b82f6"),
        ("支架 P2", (1030, 620), "#3b82f6"),
        ("锚索 A1", (1210, 330), "#22c55e"),
        ("微震 M1", (1400, 430), "#a855f7"),
        ("微震 M2", (455, 600), "#a855f7"),
    ]
    for label, (x, y), color in sensors:
        draw.ellipse((x - 20, y - 20, x + 20, y + 20), fill=color, outline="#ffffff", width=4)
        if label == "微震 M2":
            label_box(draw, (x + 28, y - 32), label)
        elif label.startswith("支架"):
            label_box(draw, (x + 28, y - 12), label)
        else:
            label_box(draw, (x + 28, y - 18), label)
    draw.rounded_rectangle((250, 835, 1550, 950), radius=20, fill="#ffffff", outline="#cbd5e1", width=2)
    draw.text((290, 865), "布设原则：顶板应力、离层、下沉、支架阻力、锚索受力、微震能量多源协同；风险阶段提高采样频率并联动三端处置。", fill="#0f172a", font=font(28))
    img.save(path)
    return path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    for name, size, color in [
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 13, "1F4E79"),
        ("Heading 3", 11.5, "1F4E79"),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(21)
    run.font.color.rgb = RGBColor.from_string("0F172A")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string("1F4E79")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("比赛提交材料 | 2026 年 8 月").italic = True
    doc.add_paragraph()


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Inches(0.28) if style is None else None
    p.add_run(text)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_props = table.rows[0]._tr.get_or_add_trPr()
    header_props.append(OxmlElement("w:tblHeader"))
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "#E8EEF5")
        set_cell_text(cell, header, True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell_text(cell, value)
    for table_row in table.rows:
        row_props = table_row._tr.get_or_add_trPr()
        row_props.append(OxmlElement("w:cantSplit"))


def add_picture(doc: Document, path: Path, caption: str, width: float = 6.3) -> None:
    if not path.exists():
        add_para(doc, f"图示素材未找到：{path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    picture = run.add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(caption).italic = True


def report_one(figs: dict[str, Path]) -> Path:
    doc = Document()
    style_doc(doc)
    add_title(doc, "基于数字孪生的煤矿顶板灾变智能预警与可视化决策系统", "总体技术方案报告")
    doc.add_heading("摘要", level=1)
    add_para(doc, "本方案面向煤矿顶板灾变智能预警场景，针对监测数据碎片化、多源融合困难、预警模型适应性不足、可视化程度低和管控响应滞后等问题，提出“多源感知、标准接入、融合预警、数字孪生、三端协同、闭环处置”的总体技术路线。系统以顶板应力、离层、下沉、支架阻力、锚杆锚索受力和微震能量等指标为核心输入，通过模型服务输出风险分值、风险等级、贡献因子和判别依据，并在三维孪生工作面中进行风险云图和灾变过程表达。")
    add_para(doc, "当前平台已接入老师提供的 20,000 行真实监测 CSV 和 XGBoost 模型，具备顶板应力场、位移场、综合风险场切换以及六阶段灾变演示能力。生产系统采用 Vercel、Render 与 Supabase 联动部署，支持企业、监管、智库、只读和超级管理员五类角色。")
    doc.add_heading("关键词", level=1)
    add_para(doc, "数字孪生；煤矿顶板；XGBoost；多源融合；RBAC；可视化决策")
    doc.add_heading("1 项目背景与建设目标", level=1)
    add_para(doc, "顶板事故是煤矿安全生产中的重要风险来源。随着采掘深度增加、巷道围岩条件复杂化和工作面推进节奏加快，单一传感器或单一阈值已难以及时、准确地识别灾变前兆。传统管理方式常依赖二维曲线、人工巡检和事后分析，存在空间定位不直观、预警依据不透明、处置闭环不完整等问题。")
    add_bullets(doc, [
        "监测数据碎片化：应力、位移、支架阻力、微震等数据分散，难以统一研判。",
        "多源信息融合困难：单一指标无法完整表达顶板灾变演化过程。",
        "预警模型适应性不足：简单阈值容易造成漏报、误报，缺少趋势和空间联动判断。",
        "可视化程度低：风险区域、设备状态和处置对象难以在同一空间中呈现。",
        "管控响应滞后：企业、监管、智库之间缺少统一事件流和反馈机制。",
    ])
    doc.add_heading("2 整体技术路线", level=1)
    add_para(doc, "系统采用七层技术路线：感知层负责采集顶板和支护状态；传输层负责井下数据汇聚；数据层完成标准化、清洗和时空对齐；模型层完成多源融合预警；孪生层完成三维场景映射；应用层提供三端业务视图；闭环层记录预警、处置、反馈和复盘结果。")
    add_picture(doc, figs["architecture"], "图 1 系统总体架构图")
    add_table(doc, ["层级", "主要内容", "比赛要求对应"], [
        ["感知层", "应力、离层、下沉、支架阻力、锚索、微震", "深地透明感知"],
        ["数据层", "标准字段、时间戳、测点编码、质量标记", "多源数据标准接口"],
        ["模型层", "综合风险分值、等级、贡献因子、判别依据", "智能预警模型"],
        ["孪生层", "三维巷道、设备、云图、测点联动", "数字孪生可视化"],
        ["应用层", "企业端、监管端、智库端、只读端、管理端", "多角色协同管控"],
        ["闭环层", "确认、处置、督办、反馈、复盘", "管控响应闭环"],
    ])
    doc.add_heading("3 系统总体架构", level=1)
    add_para(doc, "平台采用前后端分离的生产架构。Vercel 托管登录页、业务页面、Three.js 场景和 ECharts 图表，并代理 API 请求；Render 运行 Node 服务、会话、权限校验和 RoofRisk API v1；Supabase Auth 与 PostgreSQL 持久化账户、角色、权限、状态和审计日志。")
    add_table(doc, ["模块", "实现方式", "当前状态"], [
        ["三维可视化", "Three.js", "已实现井下工作面、设备、测点和风险云图"],
        ["指标态势", "ECharts + 指标卡片", "已展示应力、位移、支架阻力、微震等指标"],
        ["数据接口", "RoofRisk API v1 / 标准化 JSON", "已支持多源指标、模型解释和事件闭环字段"],
        ["模型输出", "卡尔曼预处理 + XGBoost", "20,000 条记录，独立测试准确率 99.325%"],
        ["身份权限", "Supabase Auth + RBAC", "五类角色，后端校验与审计已上线"],
    ])
    doc.add_heading("4 核心技术创新点", level=1)
    doc.add_heading("4.1 多源融合风险评估", level=2)
    add_para(doc, "方案将顶板应力、离层、下沉、支架阻力、锚杆锚索受力和微震能量统一为模型输入，并引入趋势增长项和空间联动项，使风险结果从单点阈值报警升级为多源综合研判。")
    doc.add_heading("4.2 顶板灾变全过程数字孪生表达", level=2)
    add_para(doc, "平台将顶板灾变过程拆解为正常监测、顶板压力升高、离层异常、支架阻力异常、顶板垮落预警和应急处置六个阶段，并在三维工作面中联动风险云图、设备状态、指标卡片和处置建议。")
    six = SCREEN_ROOT / "roof-six-stage-qa-current"
    for i, name in enumerate(["01-normal-stress.png", "02-pressure-stress.png", "03-separation-displacement.png"], start=2):
        add_picture(doc, six / name, f"图 {i} 顶板灾变阶段截图：{name}", 6.2)
    doc.add_heading("4.3 多角色协同与最小权限", level=2)
    add_para(doc, "企业端面向现场处置，监管端面向远程督办与归档，智库端面向模型解释与复盘，只读端提供无写操作的受限视图，超级管理员负责账户与角色管理。各角色共享同一风险事件和模型输出，写操作同时受到前端可见性和 Render 后端权限校验保护。")
    portal = SCREEN_ROOT / "portal-three-terminal-check-current"
    add_picture(doc, portal / "enterprise.png", "图 5 企业端数字孪生监控界面", 6.2)
    add_picture(doc, portal / "regulator.png", "图 6 监管端区域风险态势界面", 6.2)
    add_picture(doc, portal / "expert.png", "图 7 智库端模型解释界面", 6.2)
    doc.add_heading("5 实施方案", level=1)
    add_numbered(doc, [
        "数据与模型：完成真实 CSV 校验、卡尔曼平滑、标准化和 XGBoost 推理。",
        "平台与接口：通过 RoofRisk API v1 统一风险分值、等级、概率和特征证据。",
        "角色与权限：完成企业、监管、智库、只读和超级管理员的 RBAC。",
        "生产与交付：完成 Vercel、Render、Supabase 部署及报告、预检和离线兜底。",
    ])
    doc.add_heading("6 监测网络设计", level=1)
    add_para(doc, "监测网络围绕工作面超前支承压力影响区、巷道顶板关键断面、液压支架受载区域和微震活动区域布设。测点编码建议采用“矿井编号-工作面编号-巷道类型-测点类型-序号”的统一格式，便于数据接口、三维点位和数据库记录保持一致。")
    add_picture(doc, figs["network"], "图 8 监测网络布设示意图")
    doc.add_heading("7 监测预警指标与阈值", level=1)
    add_table(doc, ["指标", "单位", "作用", "风险含义"], [
        ["顶板应力", "MPa", "识别支承压力集中", "异常升高说明顶板受压集中"],
        ["顶板离层", "mm", "识别层间分离", "持续增大说明围岩结构失稳"],
        ["顶板下沉", "mm", "识别巷道变形", "增长过快说明支护承载不足"],
        ["支架阻力", "kN", "判断支架受载", "异常升高或突降均需关注"],
        ["锚索受力", "kN", "判断支护系统状态", "突变表示支护受力异常"],
        ["微震能量", "J", "判断岩体破裂活动", "能量集中说明破裂活动增强"],
    ])
    add_para(doc, "综合风险分值建议采用“多源指标归一化得分 + 趋势增长项 + 空间联动项 + 模型修正项”的口径。0-30 为绿色，30-50 为关注，50-70 为黄色，70-85 为橙色，85-100 为红色。正式应用时应结合矿井地质条件、历史数据和算法验证结果校准阈值。")
    doc.add_heading("8 可行性、先进性与应用价值", level=1)
    add_bullets(doc, [
        "可行性：平台采用成熟 Web 技术和标准 JSON 数据接口，已接入老师真实数据。",
        "先进性：方案从单点阈值报警升级为 XGBoost 多源融合、三维孪生、可解释预警和多角色协同。",
        "应用价值：有助于提高灾害早期识别能力，缩短响应时间，降低人工巡检压力和误判风险。",
    ])
    doc.add_heading("9 交付与扩展边界", level=1)
    add_para(doc, "当前版本已完成生产部署、真实数据接入、权限数据库和自动化测试。提交前只需按脚本完成视频录制和跨电脑检查。后续若接入实时传感器或在线模型服务，只需保持 RoofRisk API v1 字段合同，无需重写角色页面。")
    path = OUT_DIR / "01-总体技术方案报告.docx"
    doc.save(path)
    return path


def report_two(figs: dict[str, Path]) -> Path:
    doc = Document()
    style_doc(doc)
    add_title(doc, "基于数字孪生的煤矿顶板灾变智能预警与可视化决策系统", "平台系统设计方案与智能预警模型研究报告")
    doc.add_heading("摘要", level=1)
    add_para(doc, "本报告围绕多角色可视化管控平台的软件架构、Supabase 权限数据库和 XGBoost 智能预警模型展开。平台以企业、监管、智库、只读和超级管理员为角色入口，以老师真实监测数据和 RoofRisk API v1 为基础，以数字孪生为表达方式，构建预警、处置、监管、复盘和审计闭环。")
    doc.add_heading("1 平台建设定位", level=1)
    add_para(doc, "平台定位为煤矿顶板灾变智能预警与可视化决策系统，服务对象包括生产矿井、集团或安监监管部门以及高校科研智库。系统不只展示数据，还要把模型判断、风险区域、处置建议和反馈记录串联为一个可追踪流程。")
    doc.add_heading("2 多角色可视化管控平台软件架构", level=1)
    add_picture(doc, figs["architecture"], "图 1 生产部署与多角色平台架构图")
    add_table(doc, ["端口", "使用对象", "重点能力"], [
        ["企业端", "生产矿井现场", "实时监控、风险定位、处置建议、确认反馈"],
        ["监管端", "集团/安监部门", "区域总览、风险分级、预警闭环、远程督办"],
        ["智库端", "高校/科研机构", "模型解释、参数权重、历史复盘、算法服务"],
        ["只读端", "参观/审查人员", "受限查看风险和模型信息，无写操作"],
        ["超级管理员", "系统负责人", "账户创建、角色分配、状态管理、审计日志"],
    ])
    module_heading = doc.add_heading("3 功能模块设计", level=1)
    module_heading.paragraph_format.page_break_before = True
    add_table(doc, ["功能模块", "输入", "输出"], [
        ["多源数据接入", "应力、位移、支架阻力、锚索、微震、地质信息", "统一测点数据"],
        ["数据清洗与时空对齐", "不同频率和格式的数据", "标准时间序列与三维坐标映射"],
        ["智能预警模型", "多源特征、趋势窗口、空间联动指数", "风险分值、等级、贡献因子、阶段"],
        ["数字孪生可视化", "三维模型、测点、模型输出", "应力场、位移场、综合风险场"],
        ["预警处置闭环", "风险事件和处置规则", "确认、督办、反馈、复盘记录"],
    ])
    database_heading = doc.add_heading("4 数据库设计", level=1)
    database_heading.paragraph_format.page_break_before = True
    add_para(doc, "生产系统已接入 Supabase Auth 与 PostgreSQL。Auth 保存认证主体，profiles 保存显示名称、组织和账号状态，roles、permissions、user_roles 与 role_permissions 构成 RBAC，audit_logs 记录管理员操作。高频监测结果当前由确定性 JSON 提供，后续可迁移至时序表而不改变前端合同。")
    add_picture(doc, figs["database"], "图 2 Supabase RBAC 与审计关系图")
    add_table(doc, ["数据表", "核心字段", "用途"], [
        ["profiles", "id, username, display_name, organization, status", "用户资料与启停状态"],
        ["roles", "id, key, name, description", "五类系统角色"],
        ["permissions", "id, key, name", "细粒度操作权限"],
        ["user_roles", "user_id, role_id", "用户与角色关联"],
        ["role_permissions", "role_id, permission_id", "角色与权限关联"],
        ["audit_logs", "operator_id, action, target_id, details, created_at", "管理操作审计"],
    ])
    doc.add_heading("5 深地透明感知与智能预警模型研究", level=1)
    add_para(doc, "老师数据包含 20,000 行、11 列。七项数值输入为顶板离层速率、锚杆轴力增量、锚索轴力增量、支架阻力、涌水量、微震能量和距水体/岩溶体距离，另包含数据质量类别。构建过程按设备时间切段，执行卡尔曼平滑、标准化和 XGBoost 四分类推理。")
    add_para(doc, "模型独立测试准确率为 99.325%，全量回放一致率为 99.665%。展示层通过 RoofRisk API v1 获取 risk_score、risk_level、probabilities、contribution、explanation 和 actions；在线模型可在保持字段合同的情况下替换构建期推理结果。")
    add_table(doc, ["风险分值", "风险等级", "系统动作"], [
        ["0-30", "绿色", "正常监测"],
        ["30-50", "关注", "加密监测"],
        ["50-70", "黄色", "现场复核和支护检查"],
        ["70-85", "橙色", "准备停机和管控"],
        ["85-100", "红色", "停机撤人和封控"],
    ])
    doc.add_heading("6 巷道顶板控制关键因素", level=1)
    add_bullets(doc, [
        "地质因素：顶板岩性、层理结构、断层裂隙、埋深、含水和软弱夹层。",
        "采动因素：工作面推进速度、采高采宽、超前支承压力和邻近采空区影响。",
        "支护因素：锚杆锚索长度、间距、预紧力、液压支架初撑力和工作阻力。",
        "管理因素：监测设备在线率、巡检频次、预警响应时间和处置闭环完成率。",
    ])
    doc.add_heading("7 顶板控制准则及指标体系", level=1)
    add_table(doc, ["一级指标", "二级指标", "说明"], [
        ["围岩状态", "应力、离层、下沉", "反映顶板变形和受压"],
        ["支护状态", "支架阻力、锚索受力", "反映支护系统承载能力"],
        ["动力扰动", "微震能量、事件频次", "反映岩体破裂活跃度"],
        ["趋势变化", "增长率、持续时长", "反映异常演化速度"],
        ["空间联动", "邻近测点同步异常", "反映风险区域扩展"],
    ])
    doc.add_heading("8 顶板灾变判别依据", level=1)
    add_table(doc, ["阶段", "判别依据", "处置方向"], [
        ["正常监测", "多源指标稳定，无联动异常", "常规巡检和基线记录"],
        ["顶板压力升高", "顶板应力接近关注阈值，支架阻力同步上升", "提高采样频率，复核测点"],
        ["离层异常", "离层和位移增长率明显升高", "检查锚杆锚索，降低推进速度"],
        ["支架阻力异常", "支架工作阻力异常升高或突降，支护受载不均", "调整支架初撑力，现场巡查"],
        ["垮落预警", "离层、下沉、支架阻力、微震多源耦合异常", "停机撤人，封控区域"],
        ["应急处置", "危险区域已定位，处置措施执行中", "断电封控，补强支护，持续监测"],
    ])
    six = SCREEN_ROOT / "roof-six-stage-qa-current"
    for i, name in enumerate(["04-support-risk.png", "05-fall-risk.png", "06-emergency-risk.png"], start=3):
        add_picture(doc, six / name, f"图 {i} 顶板灾变阶段截图：{name}", 6.2)
    doc.add_heading("9 演示视频脚本", level=1)
    add_numbered(doc, [
        "打开默认网址，展示企业端井下数字孪生和综合风险场。",
        "切换应力场、位移场、综合风险场，说明透明感知。",
        "播放六阶段灾变过程，展示从正常监测到应急处置的演化。",
        "切换监管端，展示区域风险总览和监管闭环。",
        "切换智库端，展示模型输入、权重、判别依据和模型服务接口。",
        "展示只读端与超级管理员，说明最小权限和账户管理。",
        "总结 Vercel、Render、Supabase 和 RoofRisk API v1 的协同架构。",
    ])
    doc.add_heading("10 当前完成度与交付状态", level=1)
    add_para(doc, "当前已完成真实数据与 XGBoost 接入、数字孪生、六阶段演示、多角色 RBAC、超级管理员、Supabase 持久化、Vercel/Render 生产部署和自动化测试。提交前只需录制正式演示视频、核对比赛模板并执行一键预检与打包。")
    path = OUT_DIR / "02-平台系统设计与智能预警模型研究报告.docx"
    doc.save(path)
    return path


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    FIG_DIR.mkdir(exist_ok=True)
    figs = {
        "architecture": create_system_architecture(),
        "database": create_database_er(),
        "network": create_monitor_network(),
    }
    outputs = [report_one(figs), report_two(figs)]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
